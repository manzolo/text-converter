import io
import markdown
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from weasyprint import HTML
from typing import BinaryIO


class TextConverter:
    """Convert processed text to various output formats."""

    @staticmethod
    def to_markdown(text: str) -> str:
        """Text is already in markdown format from AI."""
        return text

    @staticmethod
    def to_html(text: str) -> str:
        """Convert markdown text to HTML."""
        # If text already contains HTML tags, return as is
        if '<html>' in text.lower() or '<body>' in text.lower():
            return text

        # Convert markdown to HTML
        html_content = markdown.markdown(text, extensions=['extra', 'codehilite'])

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
            color: #333;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
        }}
        pre {{
            background: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""

    @staticmethod
    def to_docx(text: str) -> BinaryIO:
        """Convert text to DOCX format."""
        doc = Document()

        # Set default style
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Process text line by line
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect headings (markdown style)
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def to_pdf_reportlab(text: str) -> BinaryIO:
        """Convert text to PDF using ReportLab with simple formatting."""
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)

        # Page dimensions
        width, height = letter
        margin = 0.75 * inch
        y_position = height - margin
        max_width = width - (2 * margin)

        # Font settings
        normal_font = "Helvetica"
        bold_font = "Helvetica-Bold"
        normal_size = 11
        heading_size = 14
        line_height = 14

        # Process text line by line
        lines = text.split('\n')

        for line in lines:
            line = line.strip()

            # Check if we need a new page
            if y_position < margin:
                c.showPage()
                y_position = height - margin

            if not line:
                # Empty line - add spacing
                y_position -= line_height
                continue

            # Detect and handle headings
            is_heading = False
            if line.startswith('# '):
                line = line[2:]
                c.setFont(bold_font, heading_size)
                is_heading = True
            elif line.startswith('## '):
                line = line[3:]
                c.setFont(bold_font, heading_size - 2)
                is_heading = True
            elif line.startswith('### '):
                line = line[4:]
                c.setFont(bold_font, normal_size + 1)
                is_heading = True
            else:
                c.setFont(normal_font, normal_size)

            # Word wrap long lines
            words = line.split(' ')
            current_line = []

            for word in words:
                test_line = ' '.join(current_line + [word])
                text_width = c.stringWidth(test_line, c._fontname, c._fontsize)

                if text_width <= max_width:
                    current_line.append(word)
                else:
                    # Write current line and start new one
                    if current_line:
                        c.drawString(margin, y_position, ' '.join(current_line))
                        y_position -= line_height
                        current_line = [word]

                        # Check for new page
                        if y_position < margin:
                            c.showPage()
                            y_position = height - margin
                            c.setFont(c._fontname, c._fontsize)  # Restore font

            # Write remaining text
            if current_line:
                c.drawString(margin, y_position, ' '.join(current_line))
                y_position -= line_height + (6 if is_heading else 3)

        c.save()
        buffer.seek(0)
        return buffer

    @staticmethod
    def to_pdf_weasyprint(text: str) -> BinaryIO:
        """Convert text to PDF using WeasyPrint (HTML to PDF)."""
        html_content = TextConverter.to_html(text)
        buffer = io.BytesIO()
        # WeasyPrint API: write_pdf returns bytes, or takes target parameter
        pdf_bytes = HTML(string=html_content).write_pdf()
        buffer.write(pdf_bytes)
        buffer.seek(0)
        return buffer
